import streamlit as st
import subprocess
import os
import pandas as pd
import json
import re
from streamlit_option_menu import option_menu
import base64
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from datetime import datetime

BASE_DIR = Path(__file__).resolve().parent
REPORTS_DIR = BASE_DIR / "reports"
HISTORY_DIR = BASE_DIR / "history"
IMAGES_DIR = BASE_DIR / "images"
CURRENT_DIR = BASE_DIR

# --------------------추가
def cleanup_reports():
    import shutil
    report_dir = CURRENT_DIR / "reports"
    if report_dir.exists():
        for f in report_dir.glob("*_result.txt"):
            try:
                f.unlink()
            except:
                pass
#--------------------------

def save_df_to_docx(df: pd.DataFrame, save_path, target_ip: str):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "NanumGothic"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "NanumGothic")
    style.font.size = Pt(10)

    title = doc.add_heading(
        f"{datetime.now().strftime('%Y-%m-%d')} 취약점 점검 결과",
        level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"대상 서버 : {target_ip}")
    p.paragraph_format.space_after = Pt(12)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    doc.save(str(save_path))

def load_image_base64(path: Path) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

RAPA_LOGO = load_image_base64(IMAGES_DIR / "rapa.png")
AUTOEVER_LOGO = load_image_base64(IMAGES_DIR / "hyundai_autoever.jpg")

st.set_page_config(
    page_title="Linux Security Dashboard",
    layout="wide",
    initial_sidebar_state="collapsed"
)

if "page" not in st.session_state:
    st.session_state.page = "main"

st.markdown("""
<style>
html, body {
    height: 100%;
}

.block-container {
    display: flex;
    flex-direction: column;

    padding-top: 0;
    padding-left: 0;
    padding-right: 0;
    padding-bottom: 0 !important;
    margin-bottom: 0 !important;
}

.hero-wrapper {
    width: 100%;
    margin-left: 0;
}

.hero {
    position: relative;
    width: 100%;
    min-height: 95vh;

    display: flex;
    align-items: center;
    justify-content: center;

    background:
        linear-gradient(
            to right,
            rgba(0,0,0,0.55) 0%,
            rgba(0,0,0,0.65) 40%,
            rgba(0,0,0,0.75) 100%
        ),
        url("https://images.unsplash.com/photo-1558494949-ef010cbdcc31");

    background-size: cover;
    background-position: center;
    background-repeat: no-repeat;

    background-attachment: fixed;

    transition: min-height 0.5s ease;
}

.hero-content {
    position: relative;
    z-index: 2;
    max-width: 1000px;
    text-align: center;
    color: #ffffff;
    padding: 0 24px;
}

.hero-content h1 {
    font-size: clamp(44px, 4.5vw, 72px);
    font-weight: 700;
    letter-spacing: -1px;
    margin-bottom: 16px;
}

.hero-content p {
    font-size: clamp(18px, 1.3vw, 24px);
    opacity: 0.9;
    line-height: 1.7;
}

.hero.shrink {
    min-height: 240px;
}

.hero.shrink .hero-content h1 {
    font-size: 32px;
}

.hero.sidebar-open .hero-content {
    transform: translate(calc(-50% + 160px), -50%);
}

.section {
    max-width: 1100px;
    margin: auto;
    padding: 80px 20px 120px;
}

section[data-testid="stSidebar"] {
    background-color: #f2f2f2;
}

.nav-link {
    margin: 6px 8px;
    padding: 10px 14px !important;

    font-size: 16px;
    color: #333 !important;
    border-radius: 14px !important;
}

.nav-link:hover {
    background-color: #e5e5e5 !important;
}

.nav-link.active,
.nav-link-selected {
    background-color: #dcdcdc !important;
    color: #000 !important;
    font-weight: 700 !important;
}

.nav-link i {
    font-size: 18px;
}

button[data-testid="collapsedControl"] {
    display: flex !important;
    align-items: center;
    gap: 6px;

    padding: 6px 12px !important;
    border-radius: 20px;

    background-color: #f2f2f2;
    color: #444;
    font-weight: 600;
}

button[data-testid="collapsedControl"]::after {
    content: "menu";
    font-size: 14px;
    letter-spacing: 0.5px;
}

button[data-testid="collapsedControl"]:hover {
    background-color: #e0e0e0;
}

/* ===========================
   ENTERPRISE DESIGN SYSTEM
=========================== */

body {
    background-color: #f7f9fc;
    font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
}

/* HERO CTA */
.hero-cta {
    margin-top: 30px;
    display: flex;
    gap: 16px;
    justify-content: center;
}

.cta-primary {
    background: #ffffff;
    color: #0b1220;
    padding: 12px 22px;
    border-radius: 10px;
    font-weight: 700;
    text-decoration: none;
    transition: all .2s ease;
}

.cta-primary:hover {
    transform: translateY(-3px);
    box-shadow: 0 16px 36px rgba(0,0,0,0.2);
}

.cta-outline {
    border: 1px solid rgba(255,255,255,0.6);
    color: #ffffff;
    padding: 12px 22px;
    border-radius: 10px;
    font-weight: 600;
    text-decoration: none;
    transition: all .2s ease;
}

.cta-outline:hover {
    background: rgba(255,255,255,0.1);
}

/* SECTION TITLE */
.section-title {
    font-size: 34px;
    font-weight: 700;
    margin-bottom: 14px;
}

.section-subtitle {
    font-size: 18px;
    opacity: 0.7;
    margin-bottom: 50px;
}

/* KPI STRIP */
.kpi-strip {
    display: flex;
    justify-content: space-between;
    text-align: center;
    margin-bottom: 80px;
}

.kpi-box h3 {
    font-size: 34px;
    font-weight: 700;
    margin-bottom: 6px;
}

.kpi-box p {
    font-size: 14px;
    opacity: 0.6;
}

/* FEATURE GRID */
.feature-grid {
    display: grid;
    grid-template-columns: repeat(4, 1fr);
    gap: 30px;
}

.feature-card {
    background: #ffffff;
    border-radius: 16px;
    padding: 32px;
    box-shadow: 0 12px 32px rgba(0,0,0,0.06);
    transition: all .25s ease;
}

.feature-card:hover {
    transform: translateY(-8px);
    box-shadow: 0 20px 48px rgba(0,0,0,0.1);
}

.feature-card h4 {
    font-size: 18px;
    margin-bottom: 12px;
    font-weight: 700;
    position: relative;
    padding-left: 14px;
}

.feature-card h4::before {
    content: "";
    position: absolute;
    left: 0;
    top: 4px;
    width: 4px;
    height: 18px;
    background: #2563eb;
    border-radius: 4px;
}

.feature-card p {
    font-size: 15px;
    opacity: 0.7;
    line-height: 1.6;
}

/* RESPONSIVE */
@media (max-width: 1200px) {
    .feature-grid {
        grid-template-columns: repeat(2, 1fr);
    }
}
@media (max-width: 640px) {
    .feature-grid {
        grid-template-columns: 1fr;
    }
}

.kpi-strip {
    display: flex;
    justify-content: space-between;
    text-align: center;
    margin-bottom: 80px;
    border-top: 1px solid #e5e7eb;
    border-bottom: 1px solid #e5e7eb;
    padding: 40px 0;
}

.kpi-box:not(:last-child) {
    border-right: 1px solid #e5e7eb;
}

</style>

<style>
section[data-testid="stAppViewContainer"] {
    padding-bottom: 0 !important;
    min-height: 100vh;
    display: flex;
    flex-direction: column;
}

section[data-testid="stAppViewContainer"] > .block-container {
    flex: 1;
}

section[data-testid="stMain"] {
    padding-bottom: 0 !important;
}
</style>

<style>
/* CHECK PAGE – 카드 스타일 */
.diagnosis-wrapper {
    display: flex;
    justify-content: center;
    margin-top: 0 !important;
}

.diagnosis-card {
    width: 100%;
    max-width: 720px;
    background-color: #f8f9fa;
    padding: 32px 36px;
    border-radius: 18px;
    box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05);
}

.diagnosis-title {
    text-align: center;
    font-weight: 700;
    margin-bottom: 8px;
}

.diagnosis-desc {
    text-align: center;
    opacity: 0.8;
    line-height: 1.6;
    margin-bottom: 28px;
}
</style>

<style>
.result-wrapper {
    max-width: 1200px;
    margin: 0 auto;
}

.result-wrapper [data-testid="stStatus"],
.result-wrapper [data-testid="stAlert"] {
    width: 100% !important;
    max-width: 100% !important;
}
</style>

<style>
/* download_button를 텍스트 링크처럼 */
div[data-testid="stDownloadButton"] button {
    all: unset;                /* 버튼 스타일 제거 */
    cursor: pointer;
    color: #2563eb;            /* 링크 블루 */
    font-size: 15px;
}

div[data-testid="stDownloadButton"] button:hover {
    text-decoration: underline;
}
</style>

<script>
const updateHeroState = () => {
    const hero = document.querySelector(".hero");
    if (!hero) return;

    const sidebar = document.querySelector(
        'section[data-testid="stSidebar"]'
    );
    const sidebarOpen = sidebar && sidebar.offsetWidth > 100;

    if (window.scrollY > 160 || sidebarOpen) {
        hero.classList.add("shrink");
    } else {
        hero.classList.remove("shrink");
    }

    if (sidebarOpen) {
        hero.classList.add("sidebar-open");
    } else {
        hero.classList.remove("sidebar-open");
    }
};

window.addEventListener("scroll", updateHeroState);

const observer = new MutationObserver(updateHeroState);
observer.observe(document.body, {
    attributes: true,
    childList: true,
    subtree: true
});

const resetHero = () => {
  const hero = document.querySelector(".hero");
  if (!hero) return;
  hero.classList.remove("shrink");
  hero.classList.remove("sidebar-open");
};

resetHero();

setTimeout(resetHero, 200);

setTimeout(updateHeroState, 300);

</script>
""", unsafe_allow_html=True)


# =========================================================
# Sidebar Navigation
# =========================================================
with st.sidebar:
    selected = option_menu(
        menu_title=None,
        options=["main", "점검", "기록"],
        icons=["star-fill", "shield-check", "clock-history"],
        menu_icon="list",
        default_index=0,
        styles={
            "container": {"padding": "8px"},
            "icon": {"font-size": "18px"},
            "nav-link": {
                "font-size": "16px",
                "margin": "6px",
                "border-radius": "14px",
            },
            "nav-link-selected": {
                "background-color": "#dcdcdc",
                "color": "#000",
            },
        }
    )
    page_map = {
        "main": "main",
        "점검": "check",
        "기록": "history",
    }

    st.session_state.page = page_map[selected]

# =========================================================
# MAIN / CHECK PAGE ROUTING
# =========================================================
if st.session_state.page == "main":
     #--------------추가
    cleanup_reports()
    #------------------
    
    st.markdown("""
    <div class="hero-wrapper">
        <div class="hero">
            <div class="hero-content">
                <h1>Linux Security Automation Platform</h1>
                <p>
                    KISA 주요정보통신기반시설 기술적 취약점 분석 가이드 기반<br>
                    엔터프라이즈 리눅스 보안 진단 자동화 시스템
                </p>
                <div class="hero-cta">
                    <a href="#overview" class="cta-primary">플랫폼 소개</a>
                    <a href="#features" class="cta-outline">기능 보기</a>
                </div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="section" id="overview">
        <div class="section-title">
        Enterprise-Level Security Assessment
        </div>
        <div class="section-subtitle">
        표준 기반 점검 · 자동화 · 대규모 서버 지원
        </div>
        <div class="feature-grid" id="features">
            <div class="feature-card">
                <h4>Single Server Assessment</h4>
                <p>
                IP 입력 기반 실시간 취약점 자동 진단.
                KISA 표준 항목 기반 정밀 점검 수행.
                </p>
            </div>
            <div class="feature-card">
                <h4>Bulk Server Inspection</h4>
                <p>
                CSV 업로드 기반 다수 서버 일괄 분석.
                운영 환경에 최적화된 대규모 자동 점검.
                </p>
            </div>
            <div class="feature-card">
                <h4>Automated Reporting</h4>
                <p>
                진단 결과 자동 정리 및 Word 보고서 생성.
                감사 대응 및 보안 문서화 지원.
                </p>
            </div>
            <div class="feature-card">
                <h4>CVE Intelligence Integration</h4>
                <p>
                설정 취약점 + 공개 취약점 동시 분석.
                정책 기반 진단과 실시간 위협 인텔리전스 결합.
                </p>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)


elif st.session_state.page == "check":

    # ===============================
    # 배너
    # ===============================
    st.markdown("""
    <div style="
        width: 100%;
        overflow: hidden;
        box-shadow: 0 8px 24px rgba(0,0,0,0.08);
        margin-bottom: 32px;
    ">
        <img src="https://images.unsplash.com/photo-1550751827-4bd374c3f58b"
             style="width:100%; height:220px; object-fit:cover;">
    </div>
    """, unsafe_allow_html=True)

    # ===============================
    # DIAGNOSIS CARD
    # ===============================
    st.markdown("""
    <div class="diagnosis-wrapper">
        <div class="diagnosis-card">
            <h3 class="diagnosis-title">⚙️ 진단 설정</h3>
            <div class="diagnosis-desc">
                대상 서버 정보를 입력하여 보안 점검을 실행합니다.<br>
                SSH 접속 정보를 입력하면 Ansible 기반 점검을 수행합니다.
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<div style='height:32px'></div>", unsafe_allow_html=True)

    # ===============================
    # INPUT FORM (탭 적용: 개별 입력 vs CSV 업로드)
    # ===============================
    _, center, _ = st.columns([1, 3, 1])
    with center:
        # 탭 디자인 생성
        tab1, tab2 = st.tabs(["🎯 개별 서버 진단", "📁 대량 서버 진단 (CSV)"])

        with tab1:
            target_ip = st.text_input("대상 서버 IP", placeholder="192.168.x.x", key="single_ip")
            ssh_user = st.text_input("SSH 계정", value="", key="single_user")
            ssh_pw = st.text_input("SSH 비밀번호", type="password", key="single_pw")
            uploaded_file = None # 탭1일 때는 업로드 파일 무시

        with tab2:
            st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
            uploaded_file = st.file_uploader("서버 목록 CSV 업로드 (필수: ip, user, pw)", type=["csv"], key="bulk_upload")
            if uploaded_file:
                try:
                    df_targets = pd.read_csv(uploaded_file, encoding='utf-8-sig')
                    st.dataframe(df_targets, use_container_width=True, height=150)
                except Exception as e:
                    st.error(f"CSV 읽기 실패: {e}")

        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
        start_btn = st.button("🚀 진단 시작", use_container_width=True)

        if start_btn:
            st.session_state.pop("latest_result_ip", None)
            st.session_state.pop("latest_result_df", None)

    st.markdown("<div style='height:40px'></div>", unsafe_allow_html=True)
    st.divider()

    # ===============================
    # EXECUTE DIAGNOSIS (통합 처리 로직)
    # ===============================
    _, result_center, _ = st.columns([0.3, 6, 0.3])

    if start_btn:
        inventory_path = CURRENT_DIR / "temp_inventory.ini"
        playbook_path = CURRENT_DIR / "check_playbook.yml"
        
        # 1. 대상 확인 및 인벤토리 생성
        valid_target = False
        with open(inventory_path, "w", encoding="utf-8") as f:
            f.write("[targets]\n")
            
            # CSV 파일이 업로드된 경우 (탭2)
            if uploaded_file is not None:
                for _, row in df_targets.iterrows():
                    f.write(f"{row['ip']} ansible_user={row['user']} ansible_password={row['pw']} ansible_become_password={row['pw']}\n")
                display_msg = "다중 서버"
                valid_target = True
            
            # 개별 IP가 입력된 경우 (탭1)
            elif target_ip:
                f.write(f"{target_ip} ansible_user={ssh_user} ansible_password={ssh_pw} ansible_become_password={ssh_pw}\n")
                display_msg = target_ip
                valid_target = True

        if not valid_target:
            st.error("진단 대상을 입력하거나 CSV 파일을 업로드해주세요!")
        else:
            with result_center:
                with st.status(f"🌐 {display_msg} 진단 중...", expanded=True) as status:
                    result = subprocess.run(
                        ["ansible-playbook", "-i", str(inventory_path), str(playbook_path)],
                        capture_output=True,
                        text=True
                    )

                    if result.returncode == 0:
                        status.update(label="✅ 진단 완료!", state="complete")
                        # 단일 진단일 경우 바로 결과 세션 저장
                        if uploaded_file is None:
                            st.session_state["latest_result_ip"] = target_ip
                        st.balloons()
                        st.success(f"🎉 {display_msg} 점검 성공!")
                    else:
                        status.update(label="❌ 진단 실패", state="error")
                        st.error("진단 실행 중 오류가 발생했습니다.")
                        st.code(result.stderr)



    # =====================================================
    # RESULT REPORT (넓게)
    # =====================================================
    report_dir = CURRENT_DIR / "reports"

    # 1. 진단 결과 파일 목록 확인
    if report_dir.exists():
        report_files = sorted([f.name for f in report_dir.glob("*_result.txt")])

        if report_files:
            _, result_center, _ = st.columns([0.3, 6, 0.3])
            with result_center:
                st.markdown("<div style='height:40px'></div>", unsafe_allow_html=True)
                st.markdown("### 📋 진단 결과 리포트 선택")

                # 여러 대 진단 시 선택할 수 있는 드롭다운 메뉴
                selected_file = st.selectbox(
                    "결과를 확인할 서버를 선택하세요",
                    report_files,
                    index=0,
                    help="점검이 완료된 서버의 IP 목록입니다."
                )
                st.markdown("<div style='height:40px'></div>", unsafe_allow_html=True)
                # 선택된 파일에서 IP 추출하여 세션에 저장 (기존 로직과 연동)
                recent_ip = selected_file.replace("_result.txt", "")
                st.session_state["latest_result_ip"] = recent_ip
                report_path = report_dir / selected_file

                # --- 여기서부터 기존 리포트 출력 및 저장 로직 ---
                st.markdown(
                    f"<h3>📊 {recent_ip} 진단 결과</h3>",
                    unsafe_allow_html=True
                )

                try:
                    parsed_results = []
                    with open(report_path, "r", encoding="utf-8") as f:
                        for line in f:
                            line = line.strip()
                            # JSON 형태만 파싱
                            if line.startswith("{") and line.endswith("}"):
                                data = json.loads(line)
                                parsed_results.append({
                                    "코드": data.get("code"),
                                    "중요도": data.get("severity"),
                                    "항목": data.get("item"),
                                    "상태": data.get("status"),
                                    "상세 사유": data.get("reason"),
                                })

                    if parsed_results:
                        df = pd.DataFrame(parsed_results)
                        df = df[["코드", "중요도", "항목", "상태", "상세 사유"]]

                        df = df[df["코드"].notna()]

                        df["STATUS_ORDER"] = df["상태"].apply(
                            lambda x: 0 if "취약" in str(x) else 1
                        )

                        df["U_NUM"] = df["코드"].str.extract(r'U-(\d+)').astype(int)

                        df = df.sort_values(
                            by=["STATUS_ORDER", "U_NUM"],
                            ascending=[True, True]
                        )

                        df = df.drop(columns=["STATUS_ORDER", "U_NUM"])
                        df = df.reset_index(drop=True)

                        st.session_state["latest_result_df"] = df

                        def highlight_vulnerable(row):
                            if "취약" in str(row["상태"]):
                                return ["background-color: #ffe6e1"] * len(row)
                            return [""] * len(row)

                        st.dataframe(
                            df.style
                                .apply(highlight_vulnerable, axis=1)   # 행 배경
                                .map(lambda x: "color:red; font-weight:bold;" if "취약" in str(x) else "color:green;",
                                    subset=["상태"])
                                .map(lambda x: "color:red;" if x == "상" else "color:orange;",
                                    subset=["중요도"]),
                            use_container_width=True,
                            height=420
                        )

                        # Word 저장 기능 (기존과 동일)
                        st.markdown("<div style='height:32px'></div>", unsafe_allow_html=True)
                        if st.button(f"📝 {recent_ip} 결과 Word로 보관함 저장"):
                            from datetime import datetime
                            HISTORY_DIR = CURRENT_DIR / "history"
                            HISTORY_DIR.mkdir(exist_ok=True)

                            date_str = datetime.now().strftime("%Y-%m-%d_%H%M%S")
                            docx_path = HISTORY_DIR / f"{recent_ip}_{date_str}.docx"

                            save_df_to_docx(
                                df,
                                docx_path,
                                target_ip=recent_ip
                            )

                            st.success(f"📁 {recent_ip} 리포트가 보관함에 기록되었습니다.")

                            with open(str(docx_path), "rb") as f:
                                st.download_button(
                                    label="⬇️ Word 다운로드",
                                    data=f.read(),
                                    file_name=docx_path.name,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                )

                    else:
                        st.info(f"{recent_ip} 서버의 상세 진단 결과가 비어있습니다.")

                except Exception as e:
                    st.error(f"리포트 처리 중 오류 발생: {e}") 

# =========================================================
# HISTORY PAGE
# =========================================================
elif st.session_state.page == "history":
    #--------------추가
    cleanup_reports()
    #------------------
    
    # ===============================
    # 배너
    # ===============================
    st.markdown("""
        <div style="
            width: 100%;
            overflow: hidden;
            box-shadow: 0 8px 24px rgba(0,0,0,0.08);
            margin-bottom: 32px;
        ">
            <img src="https://images.unsplash.com/photo-1550751827-4bd374c3f58b"
                style="width:100%; height:220px; object-fit:cover;">
        </div>
        """, unsafe_allow_html=True)
    
    # ===============================
    # DIAGNOSIS CARD
    # ===============================
    st.markdown("""
    <div class="diagnosis-wrapper">
        <div class="diagnosis-card">
            <h3 class="diagnosis-title">⚙️ 진단 결과</h3>
            <div class="diagnosis-desc">
                진단 결과를 출력합니다.
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.markdown("<div style='height:40px'></div>", unsafe_allow_html=True)

    _, center, _ = st.columns([1, 3, 1])

    with center:
        st.markdown(
            """
            <div style="
                font-size: 20px;
                font-weight: 500;
                margin-bottom: 12px;
            ">
                📂 진단 기록
            </div>
            """,
            unsafe_allow_html=True
        )
        st.markdown("<div style='height:10px'></div>", unsafe_allow_html=True)
        HISTORY_DIR = CURRENT_DIR / "history"
        HISTORY_DIR.mkdir(exist_ok=True)

        files = sorted(HISTORY_DIR.glob("*.docx"), reverse=True)

        if not files:
            st.info("저장된 진단 기록이 없습니다.")
        else:
            for f in files:
                # 가로 칸 나누기 (파일명/다운로드 8 : 삭제 버튼 2)
                col_file, col_del = st.columns([8, 2])
                
                with col_file:
                    with open(f, "rb") as file_data:
                        st.download_button(
                            label=f"📄 {f.name}",
                            data=file_data,
                            file_name=f.name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key=f"download_{f.name}",
                            use_container_width=True
                        )
                
                with col_del:
                    # 개별 삭제 버튼
                    if st.button("🗑️ 삭제", key=f"del_{f.name}", use_container_width=True):
                        import os
                        
                        # 1. history 폴더의 .docx 삭제
                        if f.exists():
                            f.unlink()
                        
                        # 2. reports 폴더의 연동된 .txt 삭제
                        # 파일명 규칙에 따라 매칭 (예: IP_result.docx -> IP_result.txt)
                        txt_filename = f.name.replace(".docx", ".txt")
                        txt_file = REPORTS_DIR / txt_filename
                        
                        if txt_file.exists():
                            txt_file.unlink()
                            st.success(f"{f.name} 및 리포트 삭제 완료")
                        else:
                            st.warning(f"워드 파일은 삭제되었으나, {txt_filename} 파일을 찾을 수 없습니다.")
                        
                        st.rerun() # 화면 새로고침
    # ===============================
    # FLEX SPACER (footer 밀어내기)
    # ===============================
    st.markdown(
        "<div style='flex:1'></div>",
        unsafe_allow_html=True
    )

# =========================================================
# footer
# =========================================================
st.markdown(f"""
<style>
.app-footer {{
    width: 100%;
    margin-top: auto;
    margin-bottom: 0 !important;
    padding: 12px 0;
    border-top: 1px solid #e5e5e5;
    background-color: #ffffff;
}}

.footer-inner {{
    max-width: 1100px;
    margin: auto;
    display: flex;
    justify-content: center;
    align-items: center;
    gap: 48px;
}}

.footer-inner img {{
    height: 48px;
    object-fit: contain;
    opacity: 0.9;
}}
</style>

<div class="app-footer">
    <div class="footer-inner">
        <img src="data:image/png;base64,{RAPA_LOGO}" alt="RAPA">
        <img src="data:image/png;base64,{AUTOEVER_LOGO}" alt="Hyundai AutoEver">
    </div>
</div>
""", unsafe_allow_html=True)